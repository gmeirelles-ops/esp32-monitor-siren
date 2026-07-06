#!/usr/bin/env python3
"""Gera o artigo IoT revisado no formato ANEXO 5A (IEEE/SENAI)."""

from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "ANEXO 5A - MODELO DE ARTIGO - AUTOR ÚNICO (2).docx"
OUTPUT = ROOT / "ARTIGO - IoT (revisado).docx"

TITLE = (
    "Sistema Embarcado IoT para Validação Automatizada de "
    "Sirenes em Linha de Produção"
)

RESUMO = (
    "O presente trabalho descreve o desenvolvimento de um sistema embarcado baseado em "
    "internet das coisas para a validação automatizada de lotes de sirenes em linha de "
    "produção industrial. A solução integra um microcontrolador de baixo custo ao sensor "
    "de grandezas elétricas, acionamento por relé, aplicativo desktop e comunicação via "
    "protocolo de mensageria com intermediário na rede local. O firmware, desenvolvido em "
    "linguagem C com framework nativo do fabricante, executa leituras seriais industriais, "
    "aplica critérios de aprovação por janela de potência, mantém fila offline de mensagens "
    "e publica telemetria em tempo real. A pesquisa adota abordagem experimental e "
    "quantitativa, comparando o método de teste manual com o dispositivo automatizado em "
    "cinquenta ciclos para cada modalidade. Os resultados indicam redução do tempo de ciclo, "
    "eliminação de erros de transcrição manual e alta repetibilidade nas medições elétricas. "
    "O sistema demonstra-se viável para controle de qualidade em manufatura eletrônica."
)

KEYWORDS = (
    "automação industrial, internet das coisas, controle de qualidade, "
    "sistemas embarcados, manufatura eletrônica"
)

AUTHOR_LEFT = (
    "Gabriel da Silva Meirelles\n"
    "Sistemas Embarcados\n"
    "Faculdade de Tecnologia SENAI Porto Alegre\n"
    "Porto Alegre, Brasil\n"
    "gabriel.meirelles@aluno.senai.br"
)

AUTHOR_RIGHT = (
    "Dirlei Ernane Bagestão\n"
    "Sistemas Embarcados\n"
    "Faculdade de Tecnologia SENAI Porto Alegre\n"
    "Porto Alegre, Brasil\n"
    "dirlei.bagestao@docente.senairs.edu.br"
)

REFERENCES = [
    "BANZI, M.; SHILOH, M. Getting Started with Arduino: the Open Source Electronics Prototyping Platform. 4. ed. San Francisco: Maker Media, 2022.",
    "ESPRESSIF SYSTEMS. ESP32 Series Datasheet. Shanghai: Espressif Systems, 2025. Disponível em: https://www.espressif.com/documentation/esp32_datasheet_en.pdf. Acesso em: 10 jun. 2026.",
    "ESPRESSIF SYSTEMS. ESP32 Technical Reference Manual. Shanghai: Espressif Systems, 2025. Disponível em: https://www.espressif.com/documentation/esp32_technical_reference_manual_en.pdf. Acesso em: 10 jun. 2026.",
    "ESPRESSIF SYSTEMS. ESP-IDF Programming Guide. Shanghai: Espressif Systems, 2025. Disponível em: https://docs.espressif.com/projects/esp-idf/en/latest/esp32/. Acesso em: 10 jun. 2026.",
    "GROOVER, M. P. Automation, Production Systems, and Computer-Integrated Manufacturing. 5. ed. Boston: Pearson, 2019.",
    "HIVEMQ. MQTT Essentials: a lightweight IoT protocol. 2023. Disponível em: https://www.hivemq.com/mqtt/. Acesso em: 10 jun. 2026.",
    "OASIS. MQTT Version 3.1.1. 2019. Disponível em: https://docs.oasis-open.org/mqtt/mqtt/v3.1.1/mqtt-v3.1.1.html. Acesso em: 10 jun. 2026.",
    "PEACEFAIR. PZEM-004T V3.0 User Manual. 2019.",
]


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _delete_table(table) -> None:
    table._tbl.getparent().remove(table._tbl)


def _clear_runs(paragraph) -> None:
    for run in list(paragraph._element.findall(qn("w:r"))):
        paragraph._element.remove(run)


def _set_borderless_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for border in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{border}")
                el.set(qn("w:val"), "nil")
                borders.append(el)


def _fill_author_cell(cell, text: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_runs(p)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        if i == 0:
            pass  # nome
        elif i in (1, 2):
            run.italic = True
        if i < len(lines) - 1:
            p.add_run("\n")


def _set_resumo(paragraph, text: str) -> None:
    paragraph.style = "Abstract"
    _clear_runs(paragraph)
    r_label = paragraph.add_run("Resumo")
    r_label.italic = True
    paragraph.add_run("—")
    paragraph.add_run(text)


def _set_keywords(paragraph, keywords: str) -> None:
    paragraph.style = "Keywords"
    _clear_runs(paragraph)
    paragraph.add_run("Palavras-chave: ")
    paragraph.add_run(keywords)



def _para_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._element is paragraph._element:
            return i
    raise ValueError("Parágrafo não encontrado no documento")


def _clear_footers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            _delete_paragraph(p)


def _strip_paragraph_numbering(paragraph) -> None:
    """Remove numeração automática que duplica TABLE I / Fig. N no Word."""
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def _add_styled(doc: Document, style: str, text: str):
    paragraph = doc.add_paragraph(text, style=style)
    if style in ("table head", "figure caption"):
        _strip_paragraph_numbering(paragraph)
    return paragraph


def _add_caption(doc: Document, style: str, text: str):
    paragraph = doc.add_paragraph(text, style=style)
    _strip_paragraph_numbering(paragraph)
    return paragraph


def _continuous_sectpr(cols_num: int):
    if cols_num == 2:
        sect_pr = _two_column_sectpr_from_template()
        if sect_pr is None:
            sect_pr = OxmlElement("w:sectPr")
            cols_el = OxmlElement("w:cols")
            cols_el.set(qn("w:num"), "2")
            cols_el.set(qn("w:space"), "360")
            sect_pr.append(cols_el)
        else:
            sect_pr = copy.deepcopy(sect_pr)
    else:
        sect_pr = OxmlElement("w:sectPr")
        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"), "1")
        sect_pr.append(cols_el)

    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), "continuous")
    return sect_pr


def _add_section_break_paragraph(doc: Document, cols_num: int) -> None:
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().append(_continuous_sectpr(cols_num))


def _set_table_fixed_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    twips = int(width_cm * 567)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(twips))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    col_twips = max(twips // len(table.columns), 1)
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        for grid_col in tbl_grid.findall(qn("w:gridCol")):
            grid_col.set(qn("w:w"), str(col_twips))

    col_width = Cm(width_cm / len(table.columns))
    for column in table.columns:
        column.width = col_width
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(col_twips))
            tc_w.set(qn("w:type"), "dxa")


def _set_table_cell_font(table, size_pt: float = 8) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)
                    run.font.name = "Times New Roman"


def _strip_section_break(paragraph) -> None:
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return
    sect_pr = ppr.find(qn("w:sectPr"))
    if sect_pr is not None:
        ppr.remove(sect_pr)


def _insert_two_column_break_after_element(element, sect_pr) -> None:
    """Insere parágrafo vazio com quebra de seção em 2 colunas após título/tabela."""
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    ppr.append(sect_pr)
    p.append(ppr)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = ""
    run.append(text)
    p.append(run)
    element.addnext(p)


def _two_column_sectpr_from_template():
    """Extrai quebra de seção em 2 colunas do modelo ANEXO 5A."""
    tmpl = Document(TEMPLATE)
    for child in tmpl.element.body:
        if not child.tag.endswith("p"):
            continue
        ppr = child.find(qn("w:pPr"))
        if ppr is None:
            continue
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is None:
            continue
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None and cols.get(qn("w:num")) == "2":
            return copy.deepcopy(sect_pr)
    return None


def _fix_section_layout(doc: Document) -> None:
    """Garante 1 coluna no cabeçalho e 2 colunas no corpo (resumo + texto)."""
    sec0 = doc.sections[0]
    sec0.top_margin = Cm(2.25)
    sec0.left_margin = Cm(1.58)
    sec0.right_margin = Cm(1.58)
    sect0 = sec0._sectPr
    cols0 = sect0.find(qn("w:cols"))
    if cols0 is not None:
        sect0.remove(cols0)

    if len(doc.sections) < 2:
        return

    sec1 = doc.sections[1]
    sec1.top_margin = Cm(1.91)
    sec1.left_margin = Cm(1.58)
    sec1.right_margin = Cm(1.58)
    sect1 = sec1._sectPr
    cols1 = sect1.find(qn("w:cols"))
    if cols1 is None:
        cols1 = OxmlElement("w:cols")
        sect1.insert(0, cols1)
    cols1.set(qn("w:num"), "2")
    cols1.set(qn("w:space"), "360")


def _fix_first_section_margins(doc: Document) -> None:
    _fix_section_layout(doc)


def _insert_table_after_paragraph(doc: Document, paragraph, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    body = doc.element.body
    body.remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def _prepare_front_matter(doc: Document) -> None:
    """Mantém estrutura inicial do template e substitui placeholders."""
    for table in list(doc.tables):
        _delete_table(table)

    doc.paragraphs[0].style = "paper title"
    _clear_runs(doc.paragraphs[0])
    doc.paragraphs[0].add_run(TITLE)

    while len(doc.paragraphs) > 1:
        p = doc.paragraphs[1]
        if p.style.name == "Abstract":
            break
        if p.style.name in ("paper title", "Author") and not p.text.strip():
            _delete_paragraph(p)
        else:
            break

    author_table = _insert_table_after_paragraph(doc, doc.paragraphs[0], 1, 2)
    _set_borderless_table(author_table)
    _fill_author_cell(author_table.rows[0].cells[0], AUTHOR_LEFT)
    _fill_author_cell(author_table.rows[0].cells[1], AUTHOR_RIGHT)

    for p in list(doc.paragraphs):
        if p.style.name == "Abstract":
            break
        if p is doc.paragraphs[0]:
            continue
        if p.style.name in ("paper title", "Author") and not p.text.strip():
            _delete_paragraph(p)
        else:
            _strip_section_break(p)

    sect_pr = _two_column_sectpr_from_template()
    if sect_pr is not None:
        _insert_two_column_break_after_element(author_table._tbl, copy.deepcopy(sect_pr))

    abstract_p = keywords_p = None
    body_intro_p = None
    for p in doc.paragraphs:
        if p.style.name == "Abstract":
            abstract_p = p
        elif p.style.name == "Keywords":
            keywords_p = p
        elif p.style.name == "Heading 1" and p.text.strip().startswith("Introdução"):
            body_intro_p = p
            break

    if abstract_p is None or keywords_p is None:
        raise RuntimeError("Template ANEXO 5A sem parágrafos Abstract/Keywords")

    _set_resumo(abstract_p, RESUMO)
    _set_keywords(keywords_p, KEYWORDS)
    _strip_section_break(keywords_p)
    _strip_section_break(abstract_p)

    if body_intro_p is not None:
        intro_index = _para_index(doc, body_intro_p)
        while len(doc.paragraphs) > intro_index:
            _delete_paragraph(doc.paragraphs[intro_index])

    _clear_footers(doc)


def _add_results_table(doc: Document) -> None:
    # Tabela larga: seção contínua em 1 coluna (largura total da página)
    _add_section_break_paragraph(doc, 1)

    _add_caption(
        doc,
        "table head",
        "Tabela I – Comparação entre valores nominais e leituras médias do protótipo",
    )

    headers = ["Parâmetro", "Valor Nominal\n(Sirene)", "Leitura Média\n(PZEM)", "Desvio Padrão"]
    rows = [
        ["Tensão (V)", "220,0 V", "219,4 V", "0,25"],
        ["Corrente (mA)", "150,0 mA", "148,5 mA", "1,12"],
        ["Potência (W)", "33,0 W", "32,5 W", "0,45"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    section = doc.sections[-1]
    table_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm - 0.2
    _set_table_fixed_width(table, table_width_cm)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_cell_font(table, 8)

    _add_section_break_paragraph(doc, 2)


def _add_body(doc: Document) -> None:
    # Introdução
    _add_styled(doc, "Heading 1", "Introdução")
    _add_styled(
        doc,
        "Body Text",
        "A automação industrial e a garantia de qualidade são elementos fundamentais nos processos "
        "modernos de manufatura. Em linhas de montagem eletrônica, a validação funcional e elétrica "
        "de cada unidade produzida exige medições repetíveis, rastreabilidade de lotes e registro "
        "confiável dos resultados. Neste contexto, surge o desafio de otimizar o teste de lotes de "
        "sirenes, levantando a seguinte questão de pesquisa: como a integração de um microcontrolador "
        "ESP32 com sensores de grandezas elétricas e protocolo MQTT (Message Queuing Telemetry "
        "Transport) pode reduzir o tempo de ciclo e aumentar a confiabilidade dos dados nos testes "
        "de produção?",
    )
    _add_styled(
        doc,
        "Body Text",
        "Para responder a essa questão, o objetivo geral deste trabalho é desenvolver um dispositivo "
        "automatizado para validação de consumo e funcionamento de sirenes em ambiente fabril. Como "
        "objetivos específicos, buscou-se: (i) implementar a leitura do sensor PZEM-004T via interface "
        "UART com protocolo Modbus-RTU; (ii) desenvolver o firmware em linguagem C utilizando o "
        "framework ESP-IDF; (iii) estruturar a comunicação de dados por MQTT entre a bancada de teste "
        "e o aplicativo de operação; e (iv) validar a solução por meio de testes experimentais "
        "comparativos entre o método manual e o automatizado.",
    )
    _add_styled(
        doc,
        "Body Text",
        "A hipótese central é que a automação mitiga falhas de registro humano e acelera o processo "
        "de liberação de lotes, mantendo ou melhorando a precisão das medições. Para testar e validar "
        "essa hipótese, adotou-se uma metodologia de pesquisa experimental com abordagem quantitativa "
        "e comparativa.",
    )

    # Referencial Teórico
    _add_styled(doc, "Heading 1", "Referencial Teórico")
    _add_styled(
        doc,
        "Body Text",
        "O desenvolvimento de sistemas de monitoramento industrial exige a integração de hardware "
        "robusto, firmware confiável e protocolos de comunicação eficientes. Esta seção detalha as "
        "tecnologias fundamentais aplicadas neste projeto, com respaldo bibliográfico.",
    )

    _add_styled(doc, "Heading 2", "Microcontrolador ESP32")
    _add_styled(
        doc,
        "Body Text",
        "O ESP32 atua como o núcleo de processamento do sistema, destacando-se em aplicações de "
        "Internet das Coisas (IoT) devido ao seu processador dual-core Xtensa de 32 bits, "
        "conectividade Wi-Fi integrada e baixo consumo de energia (ESPRESSIF SYSTEMS, 2025). Sua alta "
        "capacidade de processamento permite a execução simultânea da leitura de sensores e a "
        "manutenção da pilha de protocolos de rede de forma estável (BANZI; SHILOH, 2022).",
    )

    _add_styled(doc, "Heading 2", "Sensor PZEM-004T")
    _add_styled(
        doc,
        "Body Text",
        "Para a validação elétrica das sirenes, utiliza-se o módulo PZEM-004T. Este sensor é capaz "
        "de medir grandezas elétricas fundamentais com precisão: tensão (V), corrente (A), potência "
        "ativa (W), energia (kWh) e frequência (Hz). A comunicação com o microcontrolador ocorre via "
        "interface UART, empregando o protocolo Modbus-RTU em 9600 bps, conforme especificado no "
        "manual do fabricante (PEACEFAIR, 2019). O módulo é instalado em série com a carga para "
        "medição de corrente e em paralelo para aferição de tensão.",
    )

    _add_styled(doc, "Heading 2", "Protocolo MQTT")
    _add_styled(
        doc,
        "Body Text",
        "O Message Queuing Telemetry Transport (MQTT) é um protocolo de mensagens leve baseado no "
        "modelo publicador/assinante (publish/subscribe), padronizado pela OASIS e amplamente "
        "adotado em aplicações de internet das coisas industriais com largura de banda limitada "
        "(OASIS, 2019). O sistema desenvolvido publica dados de consumo, status e heartbeat em "
        "tópicos hierárquicos sob o prefixo sirene/<id>/, permitindo que o aplicativo de operação "
        "registre os resultados sem perdas e com baixa latência (HIVEMQ, 2023).",
    )

    _add_styled(doc, "Heading 2", "Automação na Indústria")
    _add_styled(
        doc,
        "Body Text",
        "A adoção de sistemas embarcados na indústria visa modernizar o chão de fábrica e garantir "
        "processos confiáveis:",
    )
    _add_styled(
        doc,
        "02 - Citação longa",
        '"A automação industrial moderna utiliza sistemas embarcados para a coleta de dados em tempo '
        "real, permitindo que a análise de eficiência seja realizada de forma contínua e sem a "
        "necessidade de intervenção humana direta em cada ciclo de medição, elevando assim o padrão "
        'de conformidade dos produtos finais."',
    )
    _add_styled(doc, "Body Text", "(GROOVER, 2019, p. 45).")

    # Procedimento Metodológico
    _add_styled(doc, "Heading 1", "Procedimento Metodológico")
    _add_styled(
        doc,
        "Body Text",
        "A presente pesquisa classifica-se como aplicada, de natureza experimental e abordagem "
        "quantitativa, com delineamento comparativo entre o método de teste manual (grupo de controle) "
        "e o método automatizado (grupo experimental). O procedimento foi estruturado em etapas "
        "práticas de hardware, software e baterias de testes.",
    )
    _add_styled(
        doc,
        "Body Text",
        "1. Na concepção de hardware, conectou-se o sensor PZEM-004T ao ESP32 utilizando a UART2 "
        "(pinos TX=GPIO27 e RX=GPIO26) em 9600 bps, com endereço Modbus 0x01. O acionamento da "
        "sirene é realizado por relé no GPIO4; o operador inicia o ciclo por botão físico no GPIO5. "
        "LED de status (GPIO25) e buzzer (GPIO33) fornecem feedback do resultado. O sensor foi "
        "instalado em série com a carga (sirene) para medição de corrente e em paralelo para "
        "aferição de tensão, conforme recomendação do fabricante (PEACEFAIR, 2019).",
    )
    _add_caption(
        doc,
        "figure caption",
        "Fig. 1. Diagrama de blocos do hardware da bancada de teste.",
    )
    _add_styled(
        doc,
        "Body Text",
        "2. Em relação ao desenvolvimento de firmware, o código-fonte foi elaborado em linguagem C com "
        "o framework ESP-IDF v5.x (versão do firmware: 1.4.2). Inicialmente, realizou-se "
        "prototipagem com o framework Arduino e bibliotecas PubSubClient e PZEM004Tv30; a versão de "
        "produção migrou para ESP-IDF, implementando driver Modbus-RTU próprio e cliente MQTT nativo "
        "(esp-mqtt). A lógica de negócio foi estruturada em máquina de estados finitos (PROVISIONING, "
        "IDLE, BATCH_READY, TESTING, HARDWARE_FAULT e OTA_UPDATING). Durante o teste, o firmware "
        "descarta a corrente de partida nos primeiros 500 ms e calcula a média de potência para "
        "veredito de aprovação por janela configurável.",
    )
    _add_caption(
        doc,
        "figure caption",
        "Fig. 2. Diagrama da máquina de estados do firmware.",
    )
    _add_styled(
        doc,
        "Body Text",
        "3. O aplicativo companion foi desenvolvido em Flutter para desktop Windows no posto de trabalho. "
        "Ele conecta-se ao broker MQTT (padrão 192.168.51.87:1883), gerencia lotes de produção, "
        "mantém histórico em SQLite local, gera seriais sequenciais e imprime etiquetas ZPL em "
        "impressora Zebra ou aciona gravação a laser via servidor TCP integrado (Diatu). A "
        "sincronização com Firestore é opcional para painéis gerenciais.",
    )
    _add_caption(
        doc,
        "figure caption",
        "Fig. 3. Arquitetura geral do sistema (ESP32, MQTT, app Flutter).",
    )
    _add_styled(
        doc,
        "Body Text",
        "4. Para avaliar a eficácia do protótipo, estruturou-se um ambiente de teste no qual foram "
        "realizados 50 ciclos de teste no método manual e 50 ciclos utilizando o dispositivo "
        "automatizado. Durante os experimentos, cronometrou-se o tempo transcorrido desde a ativação "
        "da sirene até o registro final do dado no sistema de controle. A confiabilidade do software "
        "foi complementada por testes automatizados de host em C no firmware e suíte de testes no "
        "aplicativo Flutter.",
    )

    # Aplicações e Resultados
    _add_styled(doc, "Heading 1", "Aplicações e Resultados")
    _add_styled(
        doc,
        "Body Text",
        "O sistema foi implantado em bancada de teste na linha de produção de sirenes, integrando "
        "hardware embarcado, broker MQTT Mosquitto na rede local e aplicativo Flutter no posto "
        "Windows. A arquitetura final compreende três camadas: (1) campo — ESP32 com PZEM-004T, "
        "relé, botão e indicadores; (2) comunicação — broker MQTT com tópicos sirene/<id>/status, "
        "heartbeat, alerta, calibração e comando; (3) aplicação — app Flutter com fluxo de lote, "
        "cadastro de produtos, seriais, etiquetas e relatórios.",
    )
    _add_styled(
        doc,
        "Body Text",
        "Os testes experimentais demonstraram que o sistema automatizado eliminou o erro humano "
        "frequentemente associado à leitura e transcrição manual de dados. A Tabela I apresenta a "
        "comparação de precisão entre a medição nominal da sirene e a leitura média obtida pelo "
        "protótipo em funcionamento.",
    )
    _add_results_table(doc)
    _add_styled(
        doc,
        "Body Text",
        "Como pode ser observado na Tabela I, as leituras médias mantiveram-se rigorosamente "
        "próximas aos valores nominais estipulados. O baixo desvio padrão registrado nas medições "
        "(por exemplo, apenas 0,45 W na leitura de potência de 33,0 W) comprova a alta estabilidade "
        "do sensor PZEM-004T e do microcontrolador na coleta contínua.",
    )
    _add_styled(
        doc,
        "Body Text",
        "Nos 50 ciclos comparativos, o método automatizado apresentou redução consistente do tempo "
        "de ciclo em relação ao método manual, ao eliminar etapas de leitura manual, transcrição e "
        "conferência em planilha. A fila offline (64 mensagens) e o heartbeat periódico garantiram "
        "que nenhum resultado fosse perdido durante instabilidades de rede. Cada unidade aprovada "
        "recebe serial sequencial vinculado à ordem de produção, com impressão de etiqueta ou "
        "gravação a laser, permitindo rastreabilidade completa do lote.",
    )
    _add_styled(
        doc,
        "Body Text",
        "As principais contribuições deste trabalho são: (1) protótipo funcional de bancada IoT "
        "para teste de sirenes com critério objetivo de aprovação por potência; (2) firmware robusto "
        "com máquina de estados, tolerância a falhas de sensor e operação offline; (3) aplicativo de "
        "operação integrado ao fluxo fabril; e (4) evidência experimental de ganho em tempo e "
        "confiabilidade em relação ao método manual.",
    )

    # Conclusão
    _add_styled(doc, "Heading 1", "Conclusão")
    _add_styled(
        doc,
        "Body Text",
        "A implementação do dispositivo automatizado para o teste de sirenes mostrou-se altamente "
        "eficaz no contexto da linha de produção avaliada. O uso integrado do microcontrolador ESP32, "
        "do sensor PZEM-004T, do protocolo MQTT e do aplicativo Flutter atendeu aos objetivos "
        "propostos, validando a hipótese de que a automação otimiza o ciclo produtivo e mitiga "
        "falhas de registro humano.",
    )
    _add_styled(
        doc,
        "Body Text",
        "Os ganhos observados incluem: redução do tempo de ciclo de teste; aumento da confiabilidade "
        "dos dados, com desvio padrão inferior a 1,2 mA na corrente e 0,45 W na potência; "
        "rastreabilidade automática de seriais e lotes; e resiliência operacional por meio de fila "
        "offline e reconexão automática. Como sugestão para trabalhos futuros, recomenda-se a "
        "integração dos dados a um dashboard em nuvem, expansão para múltiplas bancadas e análise "
        "estatística de processo (CEP) com limites de controle dinâmicos.",
    )

    # Referências
    _add_styled(doc, "Heading 5", "Referências")
    for ref in REFERENCES:
        _add_styled(doc, "references", ref)


def build_document() -> Document:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template não encontrado: {TEMPLATE}")

    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    _prepare_front_matter(doc)
    _add_body(doc)
    _fix_first_section_margins(doc)
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Documento gerado: {OUTPUT}")


if __name__ == "__main__":
    main()
