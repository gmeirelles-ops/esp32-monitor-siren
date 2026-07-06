#!/usr/bin/env python3
"""Audita conformidade do artigo IoT com o modelo ANEXO 5A."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_ARTICLE = ROOT / "ARTIGO - IoT (revisado).docx"

REQUIRED_STYLES = {
    "paper title",
    "Abstract",
    "Keywords",
    "Heading 1",
    "Heading 2",
    "Body Text",
    "references",
    "table head",
    "figure caption",
    "02 - Citação longa",
}

MAIN_SECTIONS = {
    "Introdução",
    "Referencial Teórico",
    "Procedimento Metodológico",
    "Aplicações e Resultados",
    "Conclusão",
}

SUBSECTIONS = {
    "Microcontrolador ESP32",
    "Sensor PZEM-004T",
    "Protocolo MQTT",
    "Automação na Indústria",
}

FORBIDDEN_HEADER_MARKERS = (
    "Recredenciamento presencial",
    "Credenciamento EAD",
    "Faculdade de Tecnologia SENAI Porto Alegre",
)

RESUMO_ACRONYMS = re.compile(
    r"\b(IoT|ESP32|MQTT|PZEM|UART|RTU|Wi-?Fi|GPIO|ZPL|OTA|CEP|SQL|TCP)\b",
    re.IGNORECASE,
)

MARGIN_LEFT_CM = 1.58
MARGIN_TOP_CM = 2.25
MARGIN_TOLERANCE_CM = 0.05


class AuditResult:
    def __init__(self) -> None:
        self.violations: list[str] = []

    def fail(self, message: str) -> None:
        self.violations.append(message)

    @property
    def ok(self) -> bool:
        return not self.violations


def _paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def _has_page_field(footer) -> bool:
    xml = footer._element.xml
    return "PAGE" in xml or "w:fldChar" in xml


def _word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))


def _extract_resumo_body(text: str) -> str:
    if text.startswith("Resumo—"):
        return text[len("Resumo—") :].strip()
    if text.startswith("Resumo —"):
        return text[len("Resumo —") :].strip()
    return text


def _find_author_table(doc: Document):
    for table in doc.tables:
        left = table.rows[0].cells[0].text
        if "Gabriel da Silva Meirelles" in left:
            return table
    return None


def audit_document(path: Path) -> AuditResult:
    result = AuditResult()
    doc = Document(path)

    # Estilos disponíveis no documento
    style_names = {s.name for s in doc.styles}
    missing_styles = REQUIRED_STYLES - style_names
    if missing_styles:
        result.fail(f"Estilos ausentes no documento: {sorted(missing_styles)}")

    # Margens da primeira seção
    sec0 = doc.sections[0]
    if abs(sec0.left_margin.cm - MARGIN_LEFT_CM) > MARGIN_TOLERANCE_CM:
        result.fail(
            f"Margem esquerda {sec0.left_margin.cm:.2f} cm "
            f"(esperado {MARGIN_LEFT_CM} ± {MARGIN_TOLERANCE_CM})"
        )
    if abs(sec0.top_margin.cm - MARGIN_TOP_CM) > MARGIN_TOLERANCE_CM:
        result.fail(
            f"Margem superior {sec0.top_margin.cm:.2f} cm "
            f"(esperado {MARGIN_TOP_CM} ± {MARGIN_TOLERANCE_CM})"
        )

    # Elementos proibidos
    for i, p in enumerate(doc.paragraphs[:6]):
        text = _paragraph_text(p)
        if any(marker in text for marker in FORBIDDEN_HEADER_MARKERS):
            result.fail(f"Parágrafo {i}: cabeçalho institucional SENAI não permitido")
    for i, p in enumerate(doc.paragraphs):
        text = _paragraph_text(p)
        if text.startswith("Abstract") or text.startswith("Keywords:"):
            result.fail(f"Parágrafo {i}: abstract/keywords em inglês não permitido")

    for section in doc.sections:
        if _has_page_field(section.footer):
            result.fail("Rodapé contém numeração de página")
            break

    # Resumo
    abstract_paras = [p for p in doc.paragraphs if p.style.name == "Abstract"]
    if not abstract_paras:
        result.fail("Parágrafo de resumo (estilo Abstract) não encontrado")
    else:
        ap = abstract_paras[0]
        text = _paragraph_text(ap)
        if not text.startswith("Resumo—"):
            result.fail("Resumo deve iniciar com 'Resumo—' (travessão sem espaços)")
        body = _extract_resumo_body(text)
        wc = _word_count(body)
        if wc < 100 or wc > 250:
            result.fail(f"Resumo com {wc} palavras (esperado 100–250)")
        if RESUMO_ACRONYMS.search(body):
            result.fail("Resumo contém siglas ou acrônimos proibidos")

    # Palavras-chave
    kw_paras = [p for p in doc.paragraphs if p.style.name == "Keywords"]
    if not kw_paras:
        result.fail("Parágrafo de palavras-chave (estilo Keywords) não encontrado")
    else:
        text = _paragraph_text(kw_paras[0])
        if not text.lower().startswith("palavras-chave:"):
            result.fail("Palavras-chave devem iniciar com 'Palavras-chave:'")
        terms_part = text.split(":", 1)[-1].strip()
        if ";" in terms_part:
            result.fail("Palavras-chave devem ser separadas por vírgula, não ponto e vírgula")
        terms = [t.strip() for t in terms_part.split(",") if t.strip()]
        if len(terms) < 3 or len(terms) > 5:
            result.fail(f"Palavras-chave: {len(terms)} termos (esperado 3–5)")
        for term in terms:
            if term[0].isupper() and not term[0].isdigit():
                result.fail(f"Palavra-chave com inicial maiúscula: {term!r}")

    # Autoria
    author_table = _find_author_table(doc)
    if author_table is None:
        result.fail("Tabela de autoria não encontrada")
    else:
        left = author_table.rows[0].cells[0].text
        right = author_table.rows[0].cells[1].text
        if "Gabriel da Silva Meirelles" not in left:
            result.fail("Coluna do autor deve conter Gabriel da Silva Meirelles")
        if "Dirlei Ernane Bagestão" not in right:
            result.fail("Coluna do orientador deve conter Dirlei Ernane Bagestão")
        if "gabriel.meirelles@aluno.senai.br" not in left:
            result.fail("E-mail do autor ausente na tabela")
        if "dirlei.bagestao@docente.senairs.edu.br" not in right:
            result.fail("E-mail do orientador ausente na tabela")

    # Títulos principais
    h1_texts = {_paragraph_text(p) for p in doc.paragraphs if p.style.name == "Heading 1"}
    for section in MAIN_SECTIONS:
        if section not in h1_texts:
            result.fail(f"Seção principal ausente ou com estilo errado: {section}")
        if section.upper() == section and len(section) > 3:
            result.fail(f"Seção em maiúsculas: {section}")

    h1_upper = [
        _paragraph_text(p)
        for p in doc.paragraphs
        if p.style.name == "Heading 1" and _paragraph_text(p) == _paragraph_text(p).upper()
    ]
    if h1_upper:
        result.fail(f"Títulos Heading 1 em maiúsculas: {h1_upper}")

    # Subseções
    h2_texts = {_paragraph_text(p) for p in doc.paragraphs if p.style.name == "Heading 2"}
    for sub in SUBSECTIONS:
        if sub not in h2_texts:
            result.fail(f"Subseção ausente ou com estilo errado: {sub}")

    # Citação longa
    long_quotes = [p for p in doc.paragraphs if p.style.name == "02 - Citação longa"]
    if not long_quotes:
        result.fail("Citação longa de Groover (estilo '02 - Citação longa') não encontrada")

    # Figuras
    fig_caps = [p for p in doc.paragraphs if p.style.name == "figure caption"]
    if len(fig_caps) < 3:
        result.fail(f"Esperadas 3 legendas figure caption, encontradas {len(fig_caps)}")
    for cap in fig_caps:
        text = _paragraph_text(cap)
        if text.lower().startswith("figura "):
            result.fail(f"Legenda deve usar 'Fig.' e não 'Figura': {text[:40]}")
        if re.search(r"Fig\.\s*\d+\.\s*Fig\.\s*\d+\.", text, re.IGNORECASE):
            result.fail(f"Legenda de figura duplicada: {text[:50]}")
        if not re.match(r"Fig\.\s*\d+\.", text):
            result.fail(f"Formato de legenda inválido: {text[:50]}")

    # Tabela
    table_heads = [p for p in doc.paragraphs if p.style.name == "table head"]
    if not table_heads:
        result.fail("Cabeçalho de tabela (estilo table head) não encontrado")
    elif len(table_heads) > 1:
        result.fail(
            f"Esperado 1 parágrafo table head (modelo: 'Tabela I – título'), "
            f"encontrados {len(table_heads)}"
        )
    else:
        head_text = _paragraph_text(table_heads[0])
        if "Tabela I" not in head_text and "TABELA I" not in head_text:
            result.fail("Tabela deve ser numerada como Tabela I")
        if "TABLE I" in head_text and "Tabela I" in head_text:
            result.fail("Cabeçalho de tabela com numeração duplicada (TABLE I + Tabela I)")

    # Referências
    h5 = [p for p in doc.paragraphs if p.style.name == "Heading 5"]
    if not h5 or _paragraph_text(h5[0]) != "Referências":
        result.fail("Seção Referências deve usar estilo Heading 5")

    ref_entries = [p for p in doc.paragraphs if p.style.name == "references"]
    if len(ref_entries) < 8:
        result.fail(f"Esperadas 8 referências, encontradas {len(ref_entries)}")

    # Corpo não deve ser todo Normal
    body_paras = [
        p
        for p in doc.paragraphs
        if p.style.name == "Body Text" and _paragraph_text(p)
    ]
    if len(body_paras) < 15:
        result.fail(f"Poucos parágrafos Body Text ({len(body_paras)}); formatação manual provável")

    return result


def main() -> int:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_ARTICLE
    if not path.exists():
        print(f"ERRO: arquivo não encontrado: {path}", file=sys.stderr)
        return 2

    result = audit_document(path)
    if result.ok:
        print(f"OK: {path.name} conforme ANEXO 5A")
        return 0

    print(f"FALHA: {path.name} — {len(result.violations)} violação(ões):", file=sys.stderr)
    for v in result.violations:
        print(f"  - {v}", file=sys.stderr)
    return 1


if __name__ == "__main__":
    sys.exit(main())
